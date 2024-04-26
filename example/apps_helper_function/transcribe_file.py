from pydub import AudioSegment
import math
import os
import logging
from openai import OpenAI
import tiktoken
from retrying import retry
from docx import Document
import tempfile
from pyannote.audio import Pipeline
import time
# Load OpenAI API key environment variable from .env file
from api.settings import env

# Initialize the OpenAI client.
client = OpenAI(
    # This is the default and can be omitted
    api_key=env("OPEN_AI_API_KEY"),
)

# Define the maximum number of retries for API calls
MAX_RETRIES = 6

HUGGINGFACE_ACCESS_TOKEN=env("HUGGINGFACE_ACCESS_TOKEN")

# Retry decorator for API calls
def retry_on_exception(exception):
    return isinstance(exception, Exception)


@retry(stop_max_attempt_number=MAX_RETRIES, retry_on_exception=retry_on_exception, wait_exponential_multiplier=1000,
       wait_exponential_max=30000)
def make_api_call(api_function, *args, **kwargs):
    try:
        logging.info(f"Attempting API call to {api_function.__name__}")
        response = api_function(*args, **kwargs)
        logging.info(f"API call to {api_function.__name__} successful")
        print(f"API call to {api_function.__name__} successful")
        return response
    except Exception as e:
        logging.error(f"API call to {api_function.__name__} failed: {e}")
        raise


# Transcribe an MP3 file using the OpenAI Whisper model.
def transcribe_with_whisper(mp3_path):
    try:
        with open(mp3_path, "rb") as audio_file:
            transcript = make_api_call(client.audio.translations.create, model="whisper-1", file=audio_file)
        return transcript.text
    except Exception as e:
        logging.error(f"Failed to transcribe with Whisper: {e}")
        return ""


# Function to count the number of tokens in a text string.
def num_tokens_from_string(string: str, encoding_name: str = "cl100k_base") -> int:
    """Returns the number of tokens in a text string."""
    encoding = tiktoken.get_encoding(encoding_name)
    num_tokens = len(encoding.encode(string))
    return num_tokens


# Convert any audio or video file to MP3 format with volume increase and normalization.
def convert_to_mp3(file_path, volume_increase_dB=5):
    if os.path.splitext(file_path)[1].lower() == '.mp3':
        return file_path
    try:
        audio = AudioSegment.from_file(file_path)
        mp3_path = os.path.splitext(file_path)[0] + '.mp3'
        audio.export(mp3_path, format='mp3')
        return mp3_path
    except Exception as e:
        logging.error(f"Failed to convert file to MP3: {e}")
        return None


# Split the MP3 file into chunks and transcribe each chunk.
def split_and_transcribe_mp3(mp3_path, max_size_bytes):
    audio = AudioSegment.from_file(mp3_path)
    total_length_ms = len(audio)
    total_length_bytes = os.path.getsize(mp3_path)
    bytes_per_ms = total_length_bytes / total_length_ms
    chunk_length_ms = math.floor(max_size_bytes / bytes_per_ms)

    transcribed_text = ""
    try:
        for start_ms in range(0, total_length_ms, chunk_length_ms):
            end_ms = min(start_ms + chunk_length_ms, total_length_ms)
            chunk_data = audio[start_ms:end_ms]
            chunk_path = f"{os.path.splitext(mp3_path)[0]}_part{start_ms // chunk_length_ms}.mp3"
            chunk_data.export(chunk_path, format="mp3")
            logging.info(f"Exported chunk {chunk_path}")

            try:
                # Transcribe the chunk and append to the transcribed text.
                transcribed_text += transcribe_with_whisper(chunk_path) + " "
            finally:
                # Remove the chunk file after transcription.
                if os.path.exists(chunk_path):
                    os.remove(chunk_path)
                    logging.info(f"Removed chunk file {chunk_path}")
    except Exception as e:
        logging.error(f"Failed during chunk transcription: {e}")
    return transcribed_text.strip()

# Define the system prompt for GPT-4.
system_prompt = ("You are a text formatter, capable of formatting unstructured text content. You will be provided "
                 "with the results of an audio transcription. It will be in a single paragraph. You need to "
                 "format it properly, add the necessary punctuation marks, capitalization and return the "
                 "transcript in a presentable way. Since this will mostly be a transcription work performed in "
                 "the industry pertaining to Market Research, you can go through it line-by-line to understand "
                 "and introduce speaker/name labels based on the context or the general overview provided. "
                 "For speaker labels, you are strictly limited to using 'Moderator' and 'Respondent'. If there "
                 "are multiple moderators or respondents interacting, label them accordingly as Moderator 1, "
                 "Moderator 2, and Respondent 1, Respondent 2, etc., if the names of the speakers are not known. "
                 "or if the names of the speakers are known, you should label them with their name followed by "
                 "(Moderator) or (Respondent) accordingly. No other speaker labels are to be used. "
                 "Do not change/modify the transcript or derive intent from it in any way. "
                 "Just beautify the provided content. That's it.")


# Function to send chunks to GPT-4 for formatting.
def format_with_gpt4(text_chunks, input_helptext):
    print('openai call enter')
    formatted_text = ""
    previous_lines = ""  # To store the last 8-10 lines from the previous response.

    for i, chunk in enumerate(text_chunks):
        # Construct the user prompt with previous context if available.
        user_prompt = f"""
        *Do NOT reduce, modify or concatenate or alter the provided transcription content in ANY way. Your job is only to beautify the text and nothing else.*

        General information about the content is (ignore if empty): "{input_helptext}"

        If this is a part of a bigger interview/fieldwork session, you will find a small section of the previous interactions below (ignore if empty):
        "{previous_lines}"

        And, the transcription content that you need to work with is (beautify and space this properly):

        "{chunk}"
        """
        try:
            # Use make_api_call with retries for the GPT-4 API call
            response = make_api_call(
                client.chat.completions.create,
                model="gpt-4-1106-preview",
                temperature=0,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
            )
            response_text = response.choices[0].message.content
            formatted_text += response_text + " "
            logging.info(f"Formatted chunk {i + 1}/{len(text_chunks)}")
            # Update the previous_lines with the last 10 lines from the response.
            previous_lines = "\n".join(response_text.splitlines()[-10:])
        except Exception as e:
            logging.error(f"Failed to format text with GPT-4: {e}")
            formatted_text += f"[Error formatting chunk {i + 1}: {e}] "

    return formatted_text.strip()

# save_transcription_to_word document
def save_transcription_to_word(formatted_transcription, input_file_path):
  try:
      doc = Document()

      # Heading for the document
      doc_heading = "Transcription Result"
      doc.add_heading(doc_heading, 0)

      doc.add_paragraph(formatted_transcription)

      # Change the filename
      word_filename = os.path.splitext(input_file_path)[0] + '_formatted_transcription.docx'
      doc.save(word_filename)

      return word_filename

  except Exception as e:
      logging.error(f"Failed to save transcription: {e}")

# Main function.
def process_audio_file(file_path, input_helptext="", delete_converted_mp3=False):
    logging.info("Transcription process started.")

    # Convert the input file to MP3 format with volume increase and normalization.
    mp3_path = convert_to_mp3(file_path)
    
    if mp3_path is None:
        logging.error("Conversion to MP3 failed, skipping transcription.")
        print("Conversion to MP3 failed, skipping transcription.")
        return ""
    
    print('File converted into mp3....')
    # Load the pre-trained speaker diarization pipeline
    pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", use_auth_token=HUGGINGFACE_ACCESS_TOKEN)
    diarization = pipeline(mp3_path)
    
    print('Got audio diarization>>>>>>>>>>>>>')

    timestamped_transcription = ""
    for turn, _, speaker in diarization.itertracks(yield_label=True):

        # Extract the audio segment for this speaker turn
        start_time, end_time = turn.start, turn.end
        audio_segment = AudioSegment.from_mp3(mp3_path)[start_time * 1000:end_time * 1000]

        # Export the segment to a temporary file and transcribe
        with tempfile.NamedTemporaryFile(delete=True, suffix='.mp3') as temp_file:
            audio_segment.export(temp_file.name, format='mp3')
            transcription = transcribe_with_whisper(temp_file.name)

        # Append to timestamped_transcription string
        timestamped_transcription += f"{speaker} [{time.strftime('%H:%M:%S', time.gmtime(start_time))} - {time.strftime('%H:%M:%S', time.gmtime(end_time))}]: {transcription}\n"

    if delete_converted_mp3 and os.path.exists(mp3_path) and os.path.splitext(file_path)[1].lower() != '.mp3':
        os.remove(mp3_path)
        logging.info("Converted MP3 file removed.")
        print("Converted MP3 file removed.")

    return timestamped_transcription

    

    